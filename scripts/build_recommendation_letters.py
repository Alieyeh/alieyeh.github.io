from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / "cvs" / "recommendation_letters"

NAVY = "17324D"
TEAL = "287A78"
MID_GREY = RGBColor(82, 92, 101)
LIGHT_GREY = "D9E0E6"


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)

    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        if edge not in kwargs:
            continue
        edge_data = kwargs.get(edge)
        tag = "w:{}".format(edge)
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        for key, value in edge_data.items():
            element.set(qn("w:{}".format(key)), str(value))


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Cm(1.55)
    section.bottom_margin = Cm(1.45)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)
    section.header_distance = Cm(0.7)
    section.footer_distance = Cm(0.7)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
    normal.font.size = Pt(10.2)
    normal.font.color.rgb = RGBColor(28, 34, 40)
    normal.paragraph_format.space_after = Pt(5.2)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    normal.paragraph_format.line_spacing = 1.04

    styles["No Spacing"].font.name = "Aptos"
    styles["No Spacing"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")


def add_header(doc, referee):
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.columns[0].width = Cm(10.5)
    table.columns[1].width = Cm(7.1)
    table.allow_autofit = False

    left, right = table.rows[0].cells
    left.width = Cm(10.5)
    right.width = Cm(7.1)
    left.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    right.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    name = left.paragraphs[0]
    name.paragraph_format.space_after = Pt(1)
    run = name.add_run(referee["name"])
    run.bold = True
    run.font.name = "Aptos Display"
    run.font.size = Pt(15)
    run.font.color.rgb = RGBColor.from_string(NAVY)

    title = left.add_paragraph()
    title.paragraph_format.space_after = Pt(0)
    run = title.add_run(referee["title"])
    run.font.size = Pt(9.7)
    run.font.color.rgb = MID_GREY

    organisation = left.add_paragraph()
    organisation.paragraph_format.space_after = Pt(0)
    run = organisation.add_run(referee["organisation"])
    run.font.size = Pt(9.7)
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(TEAL)

    contact = right.paragraphs[0]
    contact.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    contact.paragraph_format.space_after = Pt(0)
    for index, line in enumerate(referee["contact_lines"]):
        run = contact.add_run(line)
        run.font.size = Pt(8.8)
        run.font.color.rgb = MID_GREY
        if index < len(referee["contact_lines"]) - 1:
            run.add_break()

    for cell in (left, right):
        cell.margin_top = 0
        cell.margin_bottom = 0
        set_cell_border(
            cell,
            bottom={"val": "single", "sz": "12", "space": "5", "color": TEAL},
        )

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(0)
    spacer.paragraph_format.space_before = Pt(0)
    spacer.add_run("")


def add_meta(doc, recipient_lines, subject):
    date = doc.add_paragraph()
    date.paragraph_format.space_after = Pt(5)
    run = date.add_run("[Date]")
    run.italic = True
    run.font.color.rgb = MID_GREY

    if recipient_lines:
        recipient = doc.add_paragraph()
        recipient.paragraph_format.space_after = Pt(6)
        for index, line in enumerate(recipient_lines):
            run = recipient.add_run(line)
            if index == 0:
                run.bold = True
            if index < len(recipient_lines) - 1:
                run.add_break()

    subject_paragraph = doc.add_paragraph()
    subject_paragraph.paragraph_format.space_after = Pt(7)
    subject_paragraph.paragraph_format.keep_with_next = True
    run = subject_paragraph.add_run(subject)
    run.bold = True
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor.from_string(NAVY)


def add_body(doc, salutation, paragraphs):
    greeting = doc.add_paragraph()
    greeting.paragraph_format.space_after = Pt(5)
    greeting.add_run(salutation)

    for text in paragraphs:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(5.4)
        paragraph.paragraph_format.widow_control = True
        paragraph.add_run(text)


def add_signature(doc, closing, referee):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.keep_together = True

    paragraph.add_run(closing)
    paragraph.add_run().add_break()
    paragraph.add_run().add_break()
    paragraph.add_run(referee["name"]).bold = True
    paragraph.add_run().add_break()
    title_run = paragraph.add_run(referee["title"])
    title_run.font.size = Pt(9.5)
    title_run.font.color.rgb = MID_GREY
    paragraph.add_run().add_break()
    org_run = paragraph.add_run(referee["organisation"])
    org_run.font.size = Pt(9.5)
    org_run.font.color.rgb = MID_GREY


def add_footer(doc):
    footer = doc.sections[0].footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    run = paragraph.add_run("Confidential reference")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(135, 144, 151)


def build_letter(filename, referee, recipient_lines, subject, salutation, paragraphs):
    doc = Document()
    configure_document(doc)
    add_header(doc, referee)
    add_meta(doc, recipient_lines, subject)
    add_body(doc, salutation, paragraphs)
    add_signature(doc, "Yours sincerely," if salutation != "To Whom It May Concern," else "Yours faithfully,", referee)
    add_footer(doc)

    output = OUTPUT_DIR / filename
    doc.save(output)
    return output


FATEMEH = {
    "name": "Dr Fatemeh Torabi",
    "title": "Assistant Professor in Health Data Science and Senior Researcher",
    "organisation": "University of Cambridge and Dementias Platform UK",
    "contact_lines": [
        "2 Parkhouse, 40 Queen Ediths Way",
        "Cambridge CB1 8PW, United Kingdom",
        "+44 7535 804266",
        "fatemeh.torabi@ice.cam.ac.uk",
    ],
}

ELEN = {
    "name": "Dr Elen Golightly",
    "title": "Data Scientist and Data Manager",
    "organisation": "Swansea University (Dementias Platform UK)",
    "contact_lines": [
        "31 Northeron",
        "Swansea SA3 5PJ, United Kingdom",
        "07880 986164",
        "elen.golightly@swansea.ac.uk",
    ],
}


def main():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    build_letter(
        "Fatemeh_Torabi_Cambridge_PhD_Reference_Draft.docx",
        FATEMEH,
        [
            "Postgraduate Admissions Committee",
            "PhD in Medical Science (CRUK Cambridge Institute)",
            "University of Cambridge",
        ],
        (
            "Re: Alieyeh Sarabandi Moghaddam - SW49526, "
            "Generative Modelling for Foundational Discovery in Biomedicine"
        ),
        "Dear Members of the Admissions Committee,",
        [
            (
                "I am very pleased to recommend Ms Alieyeh Sarabandi Moghaddam for the "
                "PhD studentship in Generative Modelling for Foundational Discovery in "
                "Biomedicine under the supervision of Dr Hana Aliee. I know Alieyeh "
                "through my work as her research and writing supervisor within Dementias "
                "Platform UK (DPUK), where I have observed her development as a thoughtful, "
                "technically capable and increasingly independent biomedical data scientist."
            ),
            (
                "A particularly strong example of her research judgement is our work on "
                "Choosing the Right Genomic Dataset: A Five-Pillar Framework for Researchers, "
                "published by Real World Data Science with Alieyeh as first author. She brought "
                "together scientific design, genomic assay choice, cohort characteristics, "
                "governance, quality control and analytical readiness in a clear framework for "
                "researchers. This required more than technical knowledge: she had to identify "
                "the decisions that materially affect validity, weigh competing considerations, "
                "and communicate them precisely to an interdisciplinary audience. She engaged "
                "constructively with feedback and strengthened both the argument and the writing "
                "through successive revisions."
            ),
            (
                "Alieyeh's wider work shows the same combination of computational skill and "
                "scientific care. At DPUK she develops reproducible workflows for genomic, "
                "multi-omic, longitudinal and clinical data in secure research environments. "
                "Her current research includes cohort-scale polygenic risk score analysis and "
                "the investigation of dementia heterogeneity through cross-cohort harmonisation, "
                "clustering, pathway analysis and clinically meaningful validation. She is alert "
                "to missingness, batch effects, provenance and the distinction between an "
                "interesting model and a biologically defensible conclusion."
            ),
            (
                "Her background in computer engineering and health data science is well suited "
                "to this project. She is comfortable moving between data architecture, statistical "
                "reasoning and machine-learning implementation, and she has worked with "
                "high-dimensional genomics and multimodal biomedical data using Python, R, "
                "PyTorch and high-performance computing. Importantly, her interest in generative "
                "modelling is grounded in substantive research questions: how latent biological "
                "structure can be represented, how uncertainty should be handled, and how models "
                "can support interpretable and counterfactual reasoning rather than prediction alone."
            ),
            (
                "Alieyeh is now ready to progress from the rigorous application and evaluation "
                "of computational methods to developing original methodology. She is intellectually "
                "curious, candid about assumptions, persistent when a problem is technically "
                "difficult, and receptive to challenge without losing independence of thought. "
                "These qualities, together with her experience of real biomedical data and "
                "responsible research practice, give her an excellent foundation for doctoral training."
            ),
            (
                "I recommend Alieyeh strongly for this studentship. I believe she would contribute "
                "both technical depth and mature scientific judgement to the Aliee group and the "
                "CRUK Cambridge Institute, and that she has the potential to develop into an "
                "excellent independent researcher in computational biomedicine."
            ),
        ],
    )

    build_letter(
        "Fatemeh_Torabi_General_Academic_Reference_Draft.docx",
        FATEMEH,
        [],
        "Re: Recommendation for Alieyeh Sarabandi Moghaddam",
        "To Whom It May Concern,",
        [
            (
                "I am pleased to recommend Ms Alieyeh Sarabandi Moghaddam for doctoral study, "
                "research appointments and related opportunities in health data science, "
                "computational biomedicine and biomedical informatics. I have worked with Alieyeh "
                "as a research and writing supervisor within Dementias Platform UK (DPUK), and "
                "have come to know her as an intellectually curious, conscientious and technically "
                "strong researcher."
            ),
            (
                "Our most substantial writing collaboration led to Choosing the Right Genomic "
                "Dataset: A Five-Pillar Framework for Researchers, published by Real World Data "
                "Science with Alieyeh as first author. She took a broad and technically complex "
                "topic and developed it into a practical framework spanning dataset discovery, "
                "governance, assay and technology choice, cohort design, quality control, "
                "harmonisation and research readiness. Her work demonstrated an ability to "
                "synthesise evidence, identify the decisions that matter to research validity, "
                "and explain them in clear, natural language without losing technical precision."
            ),
            (
                "Alieyeh approaches feedback as part of the research process. She listens closely, "
                "tests suggestions against the evidence, and revises with purpose rather than "
                "making superficial changes. She is also willing to question assumptions and to "
                "defend a position when the scientific reasoning supports it. This balance of "
                "receptiveness and independence is one of her strongest qualities as an emerging "
                "researcher."
            ),
            (
                "Her broader work at DPUK encompasses reproducible genomic and multi-omic "
                "pipelines, cohort-scale polygenic risk score analysis, dementia subtyping, "
                "secure data processing, metadata and provenance. She combines a BSc in Computer "
                "Engineering with an MSc in Health Data Science, allowing her to move confidently "
                "between software implementation, data engineering, statistical analysis and "
                "biomedical interpretation. She is especially careful about limitations, "
                "generalisability and the governance requirements of sensitive health data."
            ),
            (
                "Alieyeh communicates well with technical and non-technical audiences, works "
                "constructively in multidisciplinary settings, and shows genuine initiative in "
                "developing projects beyond their initial specification. She is dependable, "
                "self-directed and motivated by research that is both methodologically rigorous "
                "and useful to others."
            ),
            (
                "I recommend Alieyeh with confidence. She would bring strong analytical ability, "
                "research maturity and a collaborative approach to any doctoral programme or "
                "biomedical research team, and I expect her to continue developing into an "
                "excellent independent scientist."
            ),
        ],
    )

    build_letter(
        "Elen_Golightly_General_Professional_Reference_Draft.docx",
        ELEN,
        [],
        "Re: Recommendation for Alieyeh Sarabandi Moghaddam",
        "To Whom It May Concern,",
        [
            (
                "I am delighted to recommend Ms Alieyeh Sarabandi Moghaddam for doctoral study "
                "and for research or data-science roles in computational biomedicine. As her "
                "current line manager at Dementias Platform UK (DPUK), based at Swansea University, "
                "I have direct knowledge of both the quality of her technical work and the way she "
                "contributes to a multidisciplinary research environment."
            ),
            (
                "Alieyeh works across the full lifecycle of complex biomedical data. She has "
                "developed and improved reproducible workflows for genomic, multi-omic, clinical "
                "and longitudinal datasets within secure research environments, including data "
                "ingest, validation, transformation, provisioning, metadata and provenance. One "
                "important contribution was designing and implementing the genomics extension to "
                "DPUK's data-provisioning system, adapting an imaging-focused process to support "
                "large genomic files and MinIO object storage. She approaches this kind of work "
                "with a clear understanding that reliability, documentation and traceability are "
                "as important as making a pipeline run."
            ),
            (
                "She has also built cohort-scale analytical workflows, including a reproducible "
                "polygenic risk score pipeline for dementia and related neurological and "
                "psychiatric conditions using Python, R, PLINK, SLURM, containers and workflow "
                "orchestration. Her research on dementia heterogeneity involves cross-cohort "
                "harmonisation, feature engineering, clustering, pathway analysis and validation "
                "against clinical outcomes. Across these projects she is systematic in diagnosing "
                "data problems, testing assumptions and recording decisions so that work can be "
                "understood and reused by others."
            ),
            (
                "I have also collaborated with Alieyeh on DPUK's privacy-preserving AI work. She "
                "implemented and evaluated dementia classification models using synthetic-data "
                "and homomorphic-encryption approaches, and ran a range of privacy attacks to "
                "assess residual disclosure risk. This work required her to balance model utility, "
                "privacy, computational constraints and clear reporting. Her contribution to the "
                "resulting collaborative report showed both technical versatility and responsible "
                "scientific judgement."
            ),
            (
                "Alieyeh is dependable, organised and able to take ownership of technically "
                "ambiguous work. She raises risks early, seeks input at the right points and then "
                "works independently to deliver. She communicates clearly with data managers, "
                "software specialists and researchers, and has supported documentation, knowledge "
                "sharing and the supervision of junior colleagues. Her calm persistence is "
                "particularly valuable when resolving difficult data or infrastructure issues."
            ),
            (
                "Alieyeh brings together strong engineering foundations, biomedical understanding "
                "and a sincere commitment to reproducible and responsible research. I recommend "
                "her without hesitation for opportunities that call for technical depth, research "
                "initiative and collaborative delivery. I am confident that she will make a "
                "substantial contribution wherever she continues her career."
            ),
        ],
    )


if __name__ == "__main__":
    main()
