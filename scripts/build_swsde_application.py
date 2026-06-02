from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


OUT_DIR = Path(__file__).resolve().parents[1] / "cvs"
CV_PATH = OUT_DIR / "Alieyeh_Sarabandi_Moghaddam_SWSDE_Senior_Research_Data_Engineer_CV.docx"
STATEMENT_PATH = OUT_DIR / "Alieyeh_Sarabandi_Moghaddam_SWSDE_Supporting_Statement.docx"

NAVY = RGBColor(22, 75, 93)
GREY = RGBColor(77, 91, 97)


def style_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.62)
    section.right_margin = Inches(0.62)

    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(9.5)
    normal.paragraph_format.space_after = Pt(3)
    normal.paragraph_format.line_spacing = 1.05

    for style_name, size in [("Heading 1", 12), ("Heading 2", 10.5)]:
        style = doc.styles[style_name]
        style.font.name = "Aptos"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = NAVY
        style.paragraph_format.space_before = Pt(8)
        style.paragraph_format.space_after = Pt(3)

    for list_style in ["List Bullet", "List Bullet 2"]:
        style = doc.styles[list_style]
        style.font.name = "Aptos"
        style.font.size = Pt(9.2)
        style.paragraph_format.space_after = Pt(2)
        style.paragraph_format.line_spacing = 1.03


def add_name_block(doc: Document, subtitle: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Alieyeh Sarabandi Moghaddam")
    run.bold = True
    run.font.size = Pt(17)
    run.font.color.rgb = NAVY

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(subtitle)
    run.bold = True
    run.font.size = Pt(10.5)
    run.font.color.rgb = GREY

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "Swansea, UK | alieyeh@gmail.com | linkedin.com/in/alieyeh-sarabandi-moghaddam | "
        "github.com/Alieyeh | ORCID: 0000-0003-0566-5418"
    )
    run.font.size = Pt(8.8)


def add_section(doc: Document, title: str) -> None:
    doc.add_heading(title, level=1)


def add_para(doc: Document, text: str, *, italic: bool = False, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = italic
    run.bold = bold


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_role(doc: Document, title: str, org: str, dates: str, location: str, bullets: list[str]) -> None:
    p = doc.add_paragraph()
    left = p.add_run(title)
    left.bold = True
    left.font.color.rgb = NAVY
    p.add_run(f" | {org} | {dates} | {location}")
    add_bullets(doc, bullets)


def make_cv() -> None:
    doc = Document()
    style_doc(doc)
    add_name_block(
        doc,
        "Senior Research Data Engineer | Secure Data Environments | Research-Ready Health Data",
    )

    add_section(doc, "Profile")
    add_para(
        doc,
        "Health and genomics data scientist with a computer engineering background and experience "
        "designing secure, reproducible data workflows for complex biomedical datasets in Trusted "
        "Research Environments. Current work at Dementias Platform UK covers data ingest, validation, "
        "metadata modelling, governed provisioning, integrity monitoring, documentation and researcher "
        "support for large-scale genomic, multi-omics and longitudinal cohort data."
    )
    add_para(
        doc,
        "Strong fit for research data engineering roles that require safe and resilient pipelines, "
        "research-ready data models, clear provenance, stakeholder translation and practical governance. "
        "Experienced in Python, SQL, Linux, reproducible pipeline design and cross-functional delivery "
        "across research, governance, infrastructure and clinical-facing teams."
    )

    add_section(doc, "Targeted Evidence for SWSDE")
    add_bullets(
        doc,
        [
            "Designed and maintained semi-automated ingest and provisioning workflows for sensitive biomedical datasets inside a national TRE.",
            "Built validation checks for file structure, metadata completeness, privacy readiness, provenance and downstream analysis readiness.",
            "Developed metadata frameworks, tagging systems, standardised file structures and discovery models to make heterogeneous datasets findable and usable.",
            "Implemented scheduled integrity monitoring to detect deletion, rollback, unexpected modification and version drift in shared research storage.",
            "Produced SOP-style documentation, researcher guidance and training materials for secure data handling, workflow execution and output checking.",
            "Supervised intern work on metadata discovery infrastructure and collaborated with clinicians, epidemiologists, governance specialists, researchers and engineers.",
            "Direct OMOP implementation has not been part of current production work, but experience is strongly transferable to OMOP-aligned source profiling, source-to-target mapping, validation, documentation and common-data-model delivery.",
        ],
    )

    add_section(doc, "Technical Skills")
    add_bullets(
        doc,
        [
            "Programming and engineering: Python, R, SQL, Bash, Java, C/C++, Git, Linux, CLI wrappers, package development, automated scripting.",
            "Data engineering and modelling: PostgreSQL, SQL Server, MongoDB, SQLite, MinIO object storage, ER modelling, metadata schemas, dataset tagging, provenance and audit logs.",
            "Pipelines and infrastructure: ETL-style ingest workflows, validation tooling, workflow automation, Nextflow-style design, Snakemake, SLURM, Docker, Singularity, n8n.",
            "Research data quality: QC checks, schema validation, completeness checks, file manifests, run logs, integrity monitoring, output validation and documentation standards.",
            "Health data governance: Trusted Research Environments, secure provisioning, GDPR-aware workflows, disclosure control, Five Safes principles and ONS Safe Researcher accreditation.",
            "Analysis and visualisation: pandas, scipy, scikit-learn, PyTorch, Streamlit, R Shiny, PowerBI, Plotly, matplotlib and seaborn.",
        ],
    )

    add_section(doc, "Professional Experience")
    add_role(
        doc,
        "Genomic Data Scientist",
        "Dementias Platform UK (DPUK), Swansea University",
        "2023 - Present",
        "Swansea, UK",
        [
            "Design and maintain Python-based ingest, QC, metadata harmonisation and provisioning workflows for WGS, array, methylation and multi-omics datasets within a Trusted Research Environment.",
            "Lead operational processes that connect dataset discovery, quality checking, privacy-aware review, access approvals and secure downstream analytical use.",
            "Develop validation scripts and operational checks for data quality, file structure, metadata completeness, privacy readiness and analysis readiness.",
            "Designed structured metadata frameworks, relational schema concepts and dataset discovery structures to make heterogeneous datasets searchable, comparable and reusable by researchers.",
            "Implemented automated integrity monitoring across shared research storage, producing structured outputs for triage, audit and follow-up.",
            "Develop reproducible workflows for Linux and SLURM-based secure environments using Python, R, containers and Nextflow-style orchestration.",
            "Produce SOP-style documentation, researcher guidance and training materials covering HPC usage, workflow execution, output checking and secure data handling.",
            "Collaborate with clinicians, epidemiologists, governance specialists, infrastructure engineers and researchers to translate scientific and operational requirements into robust computational workflows.",
            "Supervise an intern contributing to omics metadata curation, catalogue design and feasibility-assessment infrastructure.",
        ],
    )
    add_role(
        doc,
        "Data Scientist - Human Genetics",
        "GlaxoSmithKline (GSK), MSc Industry Placement",
        "2023",
        "UK",
        [
            "Performed region-based association analyses using UK Biobank whole-genome sequencing data to support translational human genetics and target-prioritisation research.",
            "Developed Python-based annotation and processing workflows for rare non-coding genomic regions, integrated with HPC-based genetics pipelines.",
            "Supported QC, interpretation and communication of statistical genetics outputs in collaboration with geneticists, statisticians and bioinformatics scientists.",
            "Worked in a regulated pharmaceutical research environment, aligning computational work with reproducibility, documentation and governance expectations.",
        ],
    )
    add_role(
        doc,
        "Data Scientist - Clinical Machine Learning Research",
        "Clinical outcomes prediction project",
        "2021 - 2022",
        "Remote",
        [
            "Developed machine learning models for predicting in-hospital mortality in COVID-19 patients with diabetes using structured hospital datasets.",
            "Performed feature engineering, classification modelling, model evaluation and performance analysis under ethical approval.",
            "Co-authored a peer-reviewed publication in the Journal of Diabetes and Metabolic Disorders.",
        ],
    )

    add_section(doc, "Selected Research Data Engineering Projects")
    add_para(doc, "Secure Genomics Ingest, QC and Provisioning Framework, DPUK", bold=True)
    add_bullets(
        doc,
        [
            "Designed an end-to-end framework for onboarding multi-cohort genomics and multi-omics datasets, covering QC, metadata standardisation, data readiness assessment and governed access.",
            "Built internal tooling and SOP-style documentation to reduce manual handling, improve traceability and make data more discoverable, reproducible and analysis-ready.",
        ],
    )
    add_para(doc, "Omics Metadata Atlas and Feasibility Tool, DPUK", bold=True)
    add_bullets(
        doc,
        [
            "Defined a metadata model around cohorts, datasets, experiments, file sets, quality summaries, publications, provenance and standards mappings.",
            "Specified architecture for curation, validation, database loading, API access and dashboard exploration, including PostgreSQL, Pydantic/JSON schema, FastAPI and Streamlit/Dash design.",
        ],
    )
    add_para(doc, "Research Data Integrity Monitoring System, DPUK", bold=True)
    add_bullets(
        doc,
        [
            "Designed a Python scanning system with a SQLite state store and scheduled n8n workflows to identify file deletion, modification, rollback and timestamp anomalies.",
            "Created structured JSON/CSV outputs to support operational review, triage and audit conversations.",
        ],
    )
    add_para(doc, "Multimodal Preprocessing Pipeline, Public GitHub Project", bold=True)
    add_bullets(
        doc,
        [
            "Built a reproducible Python package and CLI-style workflow for downloading, harmonising, preprocessing, validating and reporting on HAR, EEG and ECG datasets.",
            "Implemented manifests, validation reports, resource estimates, resumable execution and unit/smoke tests.",
        ],
    )

    add_section(doc, "Education")
    add_para(
        doc,
        "MSc Health Data Science (Genomics), University of Exeter, 2022 - 2023. "
        "Thesis in collaboration with GSK: using human genetic data to inform drug discovery through annotation strategies and region-based association testing."
    )
    add_para(
        doc,
        "BSc Computer Engineering, Iran University of Science and Technology, 2016 - 2021. "
        "Thesis: disease prediction using symptom-disease network modelling and probabilistic inference software."
    )

    add_section(doc, "Publications, Reports and Knowledge Exchange")
    add_bullets(
        doc,
        [
            "Sarabandi Moghaddam A., Torabi F., Squires E., Langlands K. Choosing the Right Genomic Dataset: A Five-Pillar Framework for Researchers. Real World Data Science, published 1 June 2026.",
            "Khodabakhsh P., Asadnia A., Sarabandi Moghaddam A., et al. Prediction of in-hospital mortality rate in COVID-19 patients with diabetes mellitus using machine learning methods. Journal of Diabetes and Metabolic Disorders, 2023.",
            "Contributor to DARE UK reports on safe AI in sensitive healthcare data and privacy in synthetic data, 2024.",
            "First-author conference outputs on genomics data infrastructure, PRS workflow engineering and multi-omic dementia subtyping.",
            "ONS Safe Researcher accredited; contributor to training materials on secure genomics workflows, HPC usage and output checking.",
        ],
    )

    doc.save(CV_PATH)


def make_statement() -> None:
    doc = Document()
    style_doc(doc)
    add_name_block(
        doc,
        "Supporting Statement - Senior Research Data Engineer, South West Secure Data Environment",
    )

    add_section(doc, "Motivation and Fit")
    add_para(
        doc,
        "I am applying for the Senior Research Data Engineer role because it closely matches the work I have been doing at the boundary between secure research infrastructure, data engineering, governance and applied health research. As a Genomic Data Scientist at Dementias Platform UK, I design and maintain workflows that turn complex, sensitive biomedical datasets into traceable, research-ready assets inside a Trusted Research Environment. The South West Secure Data Environment role is particularly attractive because it applies the same principles to clinical data pipelines, NHS partnerships and reusable data services that can support high-quality, data-driven healthcare research."
    )
    add_para(
        doc,
        "My background combines an MSc in Health Data Science from the University of Exeter, a BSc in Computer Engineering, current TRE-based data infrastructure work, and regulated industry experience with GSK Human Genetics. I would bring practical experience of secure data onboarding, validation, metadata modelling, documentation, stakeholder engagement and reproducible workflow development, along with a clear commitment to developing deeper expertise in OMOP-aligned data modelling for SDE use."
    )

    add_section(doc, "Secure Data Pipelines and Research-Ready Data")
    add_para(
        doc,
        "In my current role at DPUK, I design and maintain Python-based ingest, quality-control, metadata harmonisation and provisioning workflows for large-scale WGS, array, methylation and multi-omics datasets within a Trusted Research Environment. This work requires the same operational discipline needed for SDE clinical-data pipelines: understanding heterogeneous source data, defining repeatable onboarding routes, checking file and metadata completeness, preserving provenance, supporting governed access and ensuring that datasets are usable for downstream research."
    )
    add_para(
        doc,
        "I have built validation scripts and operational checks for file structure, metadata completeness, privacy readiness and downstream analysis-readiness. I have also implemented automated integrity monitoring to detect unexpected deletion, modification, rollback and version drift across shared research storage. These systems produce structured outputs for review and audit, and they reduce the risk of silent data changes undermining reproducibility. This experience maps directly to the SWSDE requirement for robust testing, validation, monitoring, automation and reliability across clinical data pipelines."
    )
    add_para(
        doc,
        "Although my current production work has not involved implementing a full OMOP ETL, it has involved many of the foundations required for common data model delivery: source profiling, metadata standardisation, schema design, controlled field definitions, documentation of assumptions, validation logic, provenance capture and researcher-facing explanation of data limitations. I would approach OMOP-aligned structures carefully and transparently: learning the national and OHDSI standards used by the SWSDE, documenting source-to-target mappings, preserving source values where required, validating transformed data and working with clinicians, data providers and researchers to ensure the resulting model remains clinically meaningful and fit for research."
    )

    add_section(doc, "Data Modelling, Documentation and Assurance")
    add_para(
        doc,
        "A significant part of my DPUK work has focused on metadata and discovery infrastructure. I have designed metadata frameworks, relational schema concepts, dataset tagging systems and discovery structures to make heterogeneous datasets searchable, comparable and usable by researchers. I have also scoped an omics metadata atlas and feasibility tool, defining a model around cohorts, datasets, experiments, file sets, quality summaries, provenance, publications and standards mappings. The planned architecture includes reproducible Python ingestion and validation, PostgreSQL-backed catalogue storage, API access and dashboard exploration."
    )
    add_para(
        doc,
        "This gives me a strong base for developing both OMOP and non-OMOP data models within the SWSDE. I understand that not every dataset can or should be forced into a single structure, and that alternative models still need clear documentation, version control, quality rules, provenance and user guidance. I am comfortable translating informal documentation, file manifests, project records and stakeholder knowledge into structured data specifications that can be implemented, reviewed and maintained."
    )
    add_para(
        doc,
        "I work in governance-aware environments and understand the importance of assurance, auditability and clear operational processes. I am ONS Safe Researcher accredited and have working knowledge of the Five Safes framework, disclosure control principles, GDPR-aware workflows and secure data provisioning. At DPUK I have produced SOP-style documentation and researcher guidance for ingest, provisioning, workflow execution, output checking and secure data handling. I would bring the same emphasis on transparent documentation, version control and defensible decision-making to SWSDE data flows."
    )

    add_section(doc, "Stakeholder Collaboration and Technical Leadership")
    add_para(
        doc,
        "The SWSDE role requires partnership working across NHS Trusts, researchers, technical teams and non-technical stakeholders. This is one of the parts of the role that most strongly matches my current experience. At DPUK I collaborate with clinicians, epidemiologists, governance specialists, infrastructure engineers, data scientists and researchers to turn scientific and operational needs into practical workflows. My work often involves clarifying what a dataset contains, what is missing, which governance constraints apply, how the data should be organised and what researchers need to know before analysis."
    )
    add_para(
        doc,
        "I have experience acting as a bridge between research users and technical delivery. I support governed provisioning processes that connect dataset discovery, quality checking, access approvals and downstream analysis. I produce clear documentation for users with different levels of technical confidence, and I have contributed to training materials on HPC usage, reproducible workflows, output checking and secure data handling. I have also supervised an intern contributing to omics metadata curation and discovery infrastructure, setting a scoped delivery plan, defining acceptance criteria and guiding the work toward usable research-software outputs."
    )
    add_para(
        doc,
        "I can work independently across multiple strands of work while keeping the practical needs of users in view. My current role includes operational delivery, tooling, documentation, researcher support and cross-team coordination. I understand the need to balance software development with assurance requirements, administrative demands and competing research timelines. I also value inclusive, patient communication: secure data infrastructure only works when researchers, data providers and technical teams can trust the process and understand each other's constraints."
    )

    add_section(doc, "Communication, Research Outputs and Wider Contribution")
    add_para(
        doc,
        "I have a strong record of communicating technical and governance-heavy material to research audiences. In June 2026 I published a Real World Data Science article, Choosing the Right Genomic Dataset: A Five-Pillar Framework for Researchers, which translates dataset discovery, access, assay choice, cohort context, QC and analysis-readiness into practical guidance for researchers. I have also contributed to DARE UK reports on safe AI in sensitive healthcare data and privacy in synthetic data, and I have first-author conference outputs on genomics data infrastructure, PRS workflow engineering and multi-omic dementia subtyping."
    )
    add_para(
        doc,
        "These outputs reflect the kind of contribution I would like to make within the SWSDE and the wider national SDE community: not only building pipelines, but also helping establish good practice, explain trade-offs, improve reusable documentation and support researchers in using sensitive health data responsibly. I am keen to engage with national Secure Data Environment and Health Data Research Service technical networks, particularly around common data models, validation, provenance, data quality and research-ready clinical datasets."
    )

    add_section(doc, "Closing")
    add_para(
        doc,
        "I would bring to this role a rare combination of secure research environment experience, practical data engineering, health data science, metadata design, documentation, training and cross-disciplinary collaboration. I do not present myself as already being an OMOP specialist; instead, I offer directly relevant SDE/TRE delivery experience and the technical foundation to become productive in OMOP-aligned work quickly and responsibly. I am motivated by the opportunity to help the South West Secure Data Environment build reliable, reusable and well-governed data services that make clinical data safer and more useful for research."
    )

    doc.save(STATEMENT_PATH)


if __name__ == "__main__":
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    make_cv()
    make_statement()
    print(CV_PATH)
    print(STATEMENT_PATH)
